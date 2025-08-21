import pytest
import os
import sys
import tempfile
import asyncio
from pathlib import Path
from unittest.mock import patch, mock_open
from docx import Document

test_dir = Path(__file__).parent
src_dir = test_dir.parent / "src"
sys.path.insert(0, str(src_dir))

from src.utils import FileProcessor

class TestFileProcessor:
    """Test cases for FileProcessor class."""
    
    @pytest.mark.asyncio
    async def test_extract_text_from_txt_file(self):
        """Test text extraction from .txt file."""
        test_content = "This is a test content\nWith multiple lines\nAnd some unicode: тест"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(test_content)
            temp_path = f.name
        
        try:
            result = await FileProcessor.extract_text_from_file(temp_path)
            assert result == test_content.strip()
        finally:
            os.unlink(temp_path)
    
    @pytest.mark.asyncio
    async def test_extract_text_from_txt_file_encoding_fallback(self):
        """Test text extraction with encoding fallback."""
        test_content = "Test content with special chars"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='cp1251') as f:
            f.write(test_content)
            temp_path = f.name
        
        try:
            # Mock the first encoding attempt to fail
            with patch('aiofiles.open', side_effect=[
                UnicodeDecodeError('utf-8', b'', 0, 1, 'test error'),
                mock_open(read_data=test_content).return_value
            ]):
                result = await FileProcessor.extract_text_from_file(temp_path)
                assert result == test_content.strip()
        finally:
            os.unlink(temp_path)
    
    @pytest.mark.asyncio
    async def test_extract_text_from_docx_file(self):
        """Test text extraction from .docx file."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph with more content")
            doc.add_paragraph("")  # Empty paragraph should be skipped
            doc.add_paragraph("Third paragraph")
            doc.save(temp_path)
            
            result = await FileProcessor.extract_text_from_file(temp_path)
            expected = "First paragraph\n\nSecond paragraph with more content\n\nThird paragraph"
            assert result == expected
        finally:
            os.unlink(temp_path)
    
    @pytest.mark.asyncio
    async def test_extract_text_unsupported_format(self):
        """Test error handling for unsupported file formats."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                await FileProcessor.extract_text_from_file(temp_path)
        finally:
            os.unlink(temp_path)
    
    @pytest.mark.asyncio
    async def test_save_text_to_file(self):
        """Test saving text content to file."""
        test_text = "This is test content\nWith multiple lines"
        filename = "test_output.txt"
        
        result_path = await FileProcessor.save_text_to_file(test_text, filename)
        
        try:
            assert result_path == f"/tmp/{filename}"
            assert os.path.exists(result_path)
            
            with open(result_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            assert content == test_text
        finally:
            if os.path.exists(result_path):
                os.unlink(result_path)
    
    def test_validate_file_valid(self):
        """Test file validation for valid files."""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
            f.write(b"test content")
            temp_path = f.name
        
        try:
            result = FileProcessor.validate_file(temp_path, 1024)
            assert result is True
        finally:
            os.unlink(temp_path)
    
    def test_validate_file_too_large(self):
        """Test file validation for files that are too large."""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
            f.write(b"test content that is longer than limit")
            temp_path = f.name
        
        try:
            result = FileProcessor.validate_file(temp_path, 10)  # Very small limit
            assert result is False
        finally:
            os.unlink(temp_path)
    
    def test_validate_file_unsupported_format(self):
        """Test file validation for unsupported formats."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(b"test content")
            temp_path = f.name
        
        try:
            result = FileProcessor.validate_file(temp_path, 1024)
            assert result is False
        finally:
            os.unlink(temp_path)
    
    def test_validate_file_nonexistent(self):
        """Test file validation for non-existent files."""
        result = FileProcessor.validate_file("/nonexistent/path/file.txt", 1024)
        assert result is False
    
    @pytest.mark.asyncio
    async def test_extract_from_docx_error_handling(self):
        """Test error handling for corrupted DOCX files."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            f.write(b"This is not a valid DOCX file")
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError, match="Failed to process DOCX file"):
                await FileProcessor.extract_text_from_file(temp_path)
        finally:
            os.unlink(temp_path)